"""
文档解析器模块

负责从不同格式的文件中提取文本内容，支持 PDF、Word、文本文件。
"""

import json
import logging
import os
from dataclasses import dataclass, asdict
from datetime import datetime
from typing import Optional

logger = logging.getLogger(__name__)


@dataclass
class ParseResult:
    """解析结果数据类"""
    content: str                    # 提取的文本内容
    success: bool                   # 是否成功
    error_message: Optional[str]    # 错误信息
    file_type: str                  # 文件类型
    extraction_method: str          # 提取方法
    timestamp: datetime             # 提取时间

    def to_json(self) -> str:
        """序列化为 JSON
        
        将 ParseResult 对象序列化为 JSON 字符串，
        datetime 对象转换为 ISO 格式字符串。
        """
        data = asdict(self)
        # 将 datetime 转换为 ISO 格式字符串
        data['timestamp'] = self.timestamp.isoformat()
        return json.dumps(data, ensure_ascii=False)

    @classmethod
    def from_json(cls, json_str: str) -> 'ParseResult':
        """从 JSON 反序列化
        
        从 JSON 字符串反序列化为 ParseResult 对象，
        ISO 格式字符串转换回 datetime 对象。
        """
        data = json.loads(json_str)
        # 将 ISO 格式字符串转换回 datetime
        data['timestamp'] = datetime.fromisoformat(data['timestamp'])
        return cls(**data)

    def has_valid_content(self) -> bool:
        """检查是否有有效内容"""
        return bool(self.content and self.content.strip())


class DocumentParser:
    """文档解析器
    
    负责从不同格式的文件中提取文本内容。
    支持的格式: PDF, Word (.docx), 文本文件 (.txt)
    """

    # 支持的文件扩展名映射
    SUPPORTED_EXTENSIONS = {
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.doc': 'docx',  # 尝试用 docx 解析器处理
        '.txt': 'txt',
        '.md': 'txt',
        '.csv': 'txt',
        '.json': 'txt',
        '.xml': 'txt',
        '.html': 'txt',
        '.htm': 'txt',
    }

    @classmethod
    def parse(cls, file_path: str) -> ParseResult:
        """根据文件类型自动选择解析方法
        
        Args:
            file_path: 文件路径
            
        Returns:
            ParseResult: 解析结果
        """
        timestamp = datetime.now()
        
        # 检查文件是否存在
        if not os.path.exists(file_path):
            return ParseResult(
                content='',
                success=False,
                error_message=f'文件不存在: {file_path}',
                file_type='unknown',
                extraction_method='none',
                timestamp=timestamp
            )
        
        # 获取文件扩展名
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        # 检查是否支持该格式
        if ext not in cls.SUPPORTED_EXTENSIONS:
            return ParseResult(
                content='',
                success=False,
                error_message=f'不支持的文件格式: {ext}',
                file_type=ext.lstrip('.') if ext else 'unknown',
                extraction_method='none',
                timestamp=timestamp
            )
        
        file_type = cls.SUPPORTED_EXTENSIONS[ext]
        
        try:
            # 根据文件类型选择解析方法
            if file_type == 'pdf':
                content = cls.parse_pdf(file_path)
                extraction_method = 'PyPDF2'
            elif file_type == 'docx':
                content = cls.parse_docx(file_path)
                extraction_method = 'python-docx'
            else:  # txt
                content = cls.parse_txt(file_path)
                extraction_method = 'text-read'
            
            # 检查内容是否为空
            if not content or not content.strip():
                return ParseResult(
                    content='',
                    success=True,
                    error_message='文档内容为空或仅包含空白字符',
                    file_type=file_type,
                    extraction_method=extraction_method,
                    timestamp=timestamp
                )
            
            return ParseResult(
                content=content,
                success=True,
                error_message=None,
                file_type=file_type,
                extraction_method=extraction_method,
                timestamp=timestamp
            )
            
        except Exception as e:
            logger.error(f'解析文件失败 {file_path}: {str(e)}')
            return ParseResult(
                content='',
                success=False,
                error_message=f'解析失败: {str(e)}',
                file_type=file_type,
                extraction_method='none',
                timestamp=timestamp
            )


    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """解析 PDF 文件
        
        使用 PyPDF2 提取 PDF 文本内容。
        
        Args:
            file_path: PDF 文件路径
            
        Returns:
            str: 提取的文本内容
            
        Raises:
            Exception: 解析失败时抛出异常
        """
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            raise ImportError('PyPDF2 未安装，请运行: pip install PyPDF2')
        
        reader = PdfReader(file_path)
        text_parts = []
        
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        
        return '\n'.join(text_parts)

    @staticmethod
    def parse_docx(file_path: str) -> str:
        """解析 Word 文档
        
        使用 python-docx 提取 Word 文档内容，包括段落和表格。
        
        Args:
            file_path: Word 文档路径
            
        Returns:
            str: 提取的文本内容
            
        Raises:
            Exception: 解析失败时抛出异常
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError('python-docx 未安装，请运行: pip install python-docx')
        
        doc = Document(file_path)
        text_parts = []
        
        # 提取段落内容
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(' | '.join(row_text))
        
        return '\n'.join(text_parts)

    @staticmethod
    def parse_txt(file_path: str) -> str:
        """解析文本文件
        
        支持多种编码格式，自动检测并使用正确的编码。
        
        Args:
            file_path: 文本文件路径
            
        Returns:
            str: 文件内容
            
        Raises:
            Exception: 读取失败时抛出异常
        """
        # 尝试的编码列表，按优先级排序
        encodings = ['utf-8', 'utf-8-sig', 'gbk', 'gb2312', 'gb18030', 'latin-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        # 如果所有编码都失败，使用 latin-1 作为最后手段（它可以读取任何字节）
        with open(file_path, 'r', encoding='latin-1') as f:
            return f.read()
